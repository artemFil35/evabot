from telegram import Update
from telegram.ext import ContextTypes
from docx import Document

async def handle_doc(update: Update, context: ContextTypes.DEFAULT_TYPE):
    document = Document()
    document.add_heading("Документ от Eva.Юрист", 0)
    document.add_paragraph("Автоматически сформировано.")
    path = "/tmp/generated.docx"
    document.save(path)
    await update.message.reply_document(open(path, "rb"), filename="document.docx")